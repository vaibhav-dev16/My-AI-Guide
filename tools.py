# tools.py
from langchain.tools import Tool
from langchain.utilities import DuckDuckGoSearchAPIWrapper
# from langchain.utilities import SerpAPIWrapper  # Optional alternative
from docx import Document
import os

def calculator_tool() -> Tool:
    """
    Calculator tool that evaluates Python expressions.
    """
    def calculate(expr: str) -> str:
        try:
            # Warning: eval can be unsafe if inputs are untrusted
            result = eval(expr)
            return str(result)
        except Exception as e:
            return f"Error: {e}"

    return Tool(
        name="Calculator",
        func=calculate,
        description="Use this tool to perform mathematical calculations. Input should be a valid Python expression."
    )

def search_tool() -> Tool:
    """
    Search tool using DuckDuckGo.
    """
    search = DuckDuckGoSearchAPIWrapper()

    def search_query(query: str) -> str:
        results = search.run(query)
        return results

    return Tool(
        name="Search",
        func=search_query,
        description="Use this tool to search the web for answers to questions. Input should be a search query."
    )

def get_tools(use_search: bool = True) -> list:
    """
    Return a list of LangChain tool objects.
    """
    tools = [calculator_tool()]
    if use_search:
        tools.append(search_tool())
    return tools

def make_notes(message: str):
    """
    Creates or updates 'important_document.docx' with the given message.

    If the file does not exist, it will be created with a title.
    If it exists, it will append the new information to the end.
    """
    file_name = "important_document.docx"

    if os.path.exists(file_name):
        # Open existing document
        doc = Document(file_name)
    else:
        # Create new document with header
        doc = Document()
        doc.add_heading("Important Document", level=1)
        doc.add_paragraph("")  # Add a blank line

    # Add new note entry
    doc.add_paragraph(f"• {message}")

    # Save the document
    doc.save(file_name)

    print(f"📝 Note added to '{file_name}': {message}")
    return f"Note added to '{file_name}'"
